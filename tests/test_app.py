from datetime import date
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest
from openpyxl import load_workbook
from docx import Document

import app as vacation_app


@pytest.fixture()
def client(tmp_path, monkeypatch):
    monkeypatch.setattr(vacation_app, "DB_PATH", tmp_path / "test.sqlite3")
    vacation_app.app.config.update(
        TESTING=True,
        SECRET_KEY="test",
        PROFILE_UPLOAD_FOLDER=str(tmp_path / "profile-images"),
        MAIL_FROM="noreply@extrahelden.de",
        MAIL_HOSTNAME="mail.extrahelden.de",
        DKIM_SELECTOR="urlaubsplaner",
        DKIM_PRIVATE_KEY=str(tmp_path / "dkim_private.pem"),
        MAIL_MAX_ATTEMPTS=8,
        MAIL_IPV4_ONLY=True,
        INITIAL_ADMIN_PASSWORD="admin",
    )
    vacation_app.init_db()
    return vacation_app.app.test_client()


def login(client, username="admin", password="admin"):
    return client.post(
        "/login",
        data={"username": username, "password": password},
        follow_redirects=True,
    )


def create_initial_member(client, username="azubi", roles=None):
    client.post(
        "/members",
        data={
            "action": "create",
            "username": username,
            "first_name": username.title(),
            "last_name": "Muster",
            "roles": roles or ["normal"],
        },
    )
    with vacation_app.db() as conn:
        row = conn.execute(
            "SELECT * FROM users WHERE username = ?", (username,)
        ).fetchone()
        return row["id"], vacation_app.reveal_initial_password(row)


def activate_member(client, username, initial_password, new_password="new-secret"):
    return client.post(
        "/initial-login",
        data={
            "username": username,
            "initial_password": initial_password,
            "new_password": new_password,
            "password_confirmation": new_password,
        },
        follow_redirects=True,
    )


def test_admin_seed_and_login(client):
    response = login(client)
    assert response.status_code == 200
    assert b"Mitglieder verwalten" in response.data
    client.get("/logout")
    login_page = client.get("/login")
    assert b"Initialer Admin" not in login_page.data
    assert b"<code>admin</code>" not in login_page.data


def test_admin_can_create_user_with_multiple_roles(client):
    login(client)
    response = client.post(
        "/members",
        data={
            "action": "create",
            "username": "max",
            "first_name": "Max",
            "last_name": "Muster",
            "roles": ["azubi", "normal", "desksharing"],
        },
        follow_redirects=True,
    )
    assert response.status_code == 200
    with vacation_app.db() as conn:
        roles = [
            row["role"]
            for row in conn.execute(
                "SELECT role FROM user_roles WHERE user_id = 2 ORDER BY role"
            )
        ]
    assert roles == ["azubi", "desksharing", "normal"]
    with vacation_app.db() as conn:
        user = conn.execute("SELECT * FROM users WHERE id = 2").fetchone()
    password = vacation_app.reveal_initial_password(user)
    assert len(password) == 8
    assert user["must_change_password"] == 1
    assert user["email"] == "max.muster@m-a-i.de"
    assert password.encode() in response.data


def test_role_based_edit_permissions(client):
    login(client)
    _, initial = create_initial_member(client, roles=["azubi"])
    client.get("/logout")
    forced = login(client, "azubi", initial)
    assert b"Initiales Passwort" in forced.data
    activate_member(client, "azubi", initial)
    login(client, "azubi", "new-secret")
    assert (
        client.post(
            "/entry",
            data={
                "user_id": 2,
                "date": date.today().isoformat(),
                "code": "BS",
                "action": "add",
            },
        ).status_code
        == 302
    )
    assert (
        client.post(
            "/entry",
            data={
                "user_id": 2,
                "date": date.today().isoformat(),
                "code": "KR",
                "action": "add",
            },
        ).status_code
        == 403
    )


def test_weekdays_and_band_ranges(client):
    login(client)
    admin = vacation_app.load_user("1")
    start, end = vacation_app.band_range(admin, 2026)
    assert (start, end) == (date(2025, 7, 2), date(2027, 7, 2))
    days = vacation_app.weekdays_between(date(2026, 1, 1), date(2026, 1, 7))
    assert [day.weekday() for day in days] == [3, 4, 0, 1, 2]


def test_bulk_entry_overwrites_and_deletes_selection(client):
    login(client)
    today = date.today().isoformat()
    add = client.post(
        "/bulk-entry",
        json={"code": "UB", "cells": [{"user_id": 1, "date": today}]},
    )
    assert add.status_code == 200
    overwrite = client.post(
        "/bulk-entry",
        json={"code": "UG", "cells": [{"user_id": 1, "date": today}]},
    )
    assert overwrite.json["updated"] == 1
    with vacation_app.db() as conn:
        codes = [
            row["code"]
            for row in conn.execute(
                "SELECT code FROM entries WHERE user_id = 1 AND entry_date = ?",
                (today,),
            )
        ]
    assert codes == ["UG"]
    delete = client.post(
        "/bulk-entry", json={"delete": True, "cells": [{"user_id": 1, "date": today}]}
    )
    assert delete.status_code == 200
    with vacation_app.db() as conn:
        count = conn.execute(
            "SELECT COUNT(*) FROM entries WHERE user_id = 1 AND entry_date = ?",
            (today,),
        ).fetchone()[0]
    assert count == 0


def test_import_ignores_excel_labels(client, tmp_path):
    login(client)
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Monat", "Juli"])
    ws.append(["KW", 30])
    ws.append(["Ferienzeit", None])
    ws.append(["Name", date.today()])
    ws.append(["Geplant oder Beantragt", None])
    ws.append(["Erika Mustermann", "UrlbGplntOdrBntrgt"])
    ws.append(["Felix Feiertag", "Frtg"])
    path = tmp_path / "import.xlsx"
    wb.save(path)
    with client:
        vacation_app.import_excel(path)
    with vacation_app.db() as conn:
        names = [
            row["username"]
            for row in conn.execute("SELECT username FROM users ORDER BY id")
        ]
    assert "monat" not in names
    assert "geplant.oder.beantragt" not in names
    assert "erika.mustermann" in names
    assert "felix.feiertag" in names
    with vacation_app.db() as conn:
        imported_codes = [
            row["code"]
            for row in conn.execute(
                "SELECT e.code FROM entries e JOIN users u ON u.id = e.user_id WHERE u.username = ?",
                ("erika.mustermann",),
            )
        ]
        feiertag_count = conn.execute(
            "SELECT COUNT(*) FROM entries e JOIN users u ON u.id = e.user_id WHERE u.username = ?",
            ("felix.feiertag",),
        ).fetchone()[0]
    assert imported_codes == ["UB"]
    assert feiertag_count == 0


def test_admin_can_upload_macro_enabled_excel(client):
    login(client)
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Name", date.today()])
    sheet.append(["Mia Macro", "UrlbGnhmgt"])
    stream = BytesIO()
    workbook.save(stream)
    stream.seek(0)
    response = client.post(
        "/upload",
        data={"file": (stream, "historie.xlsm")},
        content_type="multipart/form-data",
        follow_redirects=True,
    )
    assert response.status_code == 200
    with vacation_app.db() as conn:
        imported = conn.execute(
            """SELECT e.code FROM entries e JOIN users u ON u.id = e.user_id
               WHERE u.username = ?""",
            ("mia.macro",),
        ).fetchone()
    assert imported["code"] == "UG"


def test_initial_link_and_password_reset(client):
    login(client)
    user_id, initial = create_initial_member(client, "lina", ["normal"])
    link = vacation_app.initial_link("lina", initial)
    client.get("/logout")
    response = client.get(link.replace("https://urlaub.extrahelden.de", ""))
    assert b'value="lina"' in response.data
    assert initial.encode() in response.data

    login(client)
    response = client.post(
        "/members",
        data={"action": "reset-password", "user_id": user_id},
    )
    assert b"Neu erstellte Initialdatenanzeige" in response.data
    blocked = client.get("/", follow_redirects=True)
    assert b"Neu erstellte Initialdatenanzeige" in blocked.data
    acknowledged = client.post(
        "/members/acknowledge-initial-data", json={"user_id": user_id}
    )
    assert acknowledged.json == {"ok": True}
    assert client.get("/").status_code == 200
    with vacation_app.db() as conn:
        reset_user = conn.execute(
            "SELECT * FROM users WHERE id = ?", (user_id,)
        ).fetchone()
    assert vacation_app.reveal_initial_password(reset_user) != initial


def test_bulk_roles_and_ausbilder_may_only_reset_azubi(client):
    login(client)
    azubi_id, _ = create_initial_member(client, "lea", ["azubi"])
    normal_id, _ = create_initial_member(client, "tom", ["normal"])
    trainer_id, trainer_initial = create_initial_member(
        client, "trainer", ["ausbilder"]
    )
    client.post(
        "/members",
        data={
            "action": "roles-bulk",
            "roles_1": ["admin"],
            f"roles_{azubi_id}": ["azubi", "normal"],
            f"roles_{normal_id}": ["putzchef"],
            f"roles_{trainer_id}": ["ausbilder"],
        },
    )
    with vacation_app.db() as conn:
        assert vacation_app.roles_for_users(conn)[normal_id] == ["putzchef"]
    client.get("/logout")
    activate_member(client, "trainer", trainer_initial)
    login(client, "trainer", "new-secret")
    assert (
        client.post(
            "/members", data={"action": "reset-password", "user_id": azubi_id}
        ).status_code
        == 200
    )
    assert (
        client.post(
            "/members", data={"action": "reset-password", "user_id": normal_id}
        ).status_code
        == 403
    )


def test_profile_email_is_generated_and_immutable(client):
    login(client)
    profile = client.get("/profile")
    assert b"admin.benutzer@m-a-i.de" in profile.data
    saved = client.post(
        "/profile",
        data={"email": "admin@example.de", "birth_date": "1990-04-12"},
        follow_redirects=True,
    )
    assert b"admin.benutzer@m-a-i.de" in saved.data
    client.post("/profile", data={"email": "changed@example.de", "birth_date": ""})
    with vacation_app.db() as conn:
        row = conn.execute(
            "SELECT email, birth_date FROM users WHERE id = 1"
        ).fetchone()
    assert row["email"] == "admin.benutzer@m-a-i.de"
    assert row["birth_date"] is None


def test_init_db_backfills_only_missing_company_emails(client):
    with vacation_app.db() as conn:
        first_id = vacation_app.create_user(
            conn, "missing.mail", "Mia", "Muster Frau", "password", ["normal"]
        )
        second_id = vacation_app.create_user(
            conn, "manual.mail", "Manu", "Adresse", "password", ["normal"]
        )
        conn.execute("UPDATE users SET email = NULL WHERE id = ?", (first_id,))
        conn.execute(
            "UPDATE users SET email = 'personal@example.org' WHERE id = ?", (second_id,)
        )
    vacation_app.init_db()
    with vacation_app.db() as conn:
        rows = {
            row["id"]: row["email"]
            for row in conn.execute(
                "SELECT id, email FROM users WHERE id IN (?, ?)", (first_id, second_id)
            )
        }
    assert rows[first_id] == "mia.muster.frau@m-a-i.de"
    assert rows[second_id] == "personal@example.org"


def test_initial_credentials_are_dkim_signed_and_delivered_directly(
    client, monkeypatch
):
    sent = {}

    class FakeSmtp:
        def __init__(self, host, port, local_hostname, timeout):
            sent.update(
                host=host,
                port=port,
                local_hostname=local_hostname,
                timeout=timeout,
            )

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return None

        def ehlo(self, hostname):
            sent.setdefault("ehlo", []).append(hostname)

        def has_extn(self, extension):
            return extension == "starttls"

        def starttls(self, context):
            sent["starttls"] = True

        def sendmail(self, from_addr, recipients, message):
            sent.update(message=message, from_addr=from_addr, recipients=recipients)

    class MxAnswer:
        preference = 10
        exchange = "mx.m-a-i.de."

    class AAnswer:
        address = "192.0.2.25"

    monkeypatch.setattr(vacation_app.smtplib, "SMTP", FakeSmtp)
    monkeypatch.setattr(
        vacation_app.dns.resolver,
        "resolve",
        lambda _domain, kind: [MxAnswer()] if kind == "MX" else [AAnswer()],
    )
    login(client)
    create_initial_member(client, username="mail.user")
    with vacation_app.db() as conn:
        queued = conn.execute("SELECT * FROM mail_outbox").fetchone()
    assert queued["recipient"] == "mail.user.muster@m-a-i.de"
    assert queued["message"].startswith(b"DKIM-Signature:")
    assert b"https://urlaub.extrahelden.de/initial-login?token=" in queued["message"]

    assert vacation_app.deliver_outbox_once() is True

    assert sent["starttls"] is True
    assert sent["host"] == "192.0.2.25"
    assert sent["port"] == 25
    assert sent["local_hostname"] == "mail.extrahelden.de"
    assert sent["from_addr"] == "noreply@extrahelden.de"
    assert sent["recipients"] == ["mail.user.muster@m-a-i.de"]
    with vacation_app.db() as conn:
        delivered = conn.execute(
            "SELECT status, attempts FROM mail_outbox WHERE id = ?", (queued["id"],)
        ).fetchone()
    assert delivered["status"] == "delivered"
    assert delivered["attempts"] == 0


def test_permanent_smtp_rejection_is_not_retried(client, monkeypatch):
    class RejectingSmtp:
        def __init__(self, *_args, **_kwargs):
            pass

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return None

        def ehlo(self, _hostname):
            pass

        def has_extn(self, _extension):
            return False

        def sendmail(self, _from_addr, recipients, _message):
            raise vacation_app.smtplib.SMTPRecipientsRefused(
                {recipients[0]: (550, b"5.5.0 mailbox unavailable")}
            )

    class MxAnswer:
        preference = 10
        exchange = "mx.example.net."

    class AAnswer:
        address = "192.0.2.30"

    monkeypatch.setattr(vacation_app.smtplib, "SMTP", RejectingSmtp)
    monkeypatch.setattr(
        vacation_app.dns.resolver,
        "resolve",
        lambda _domain, kind: [MxAnswer()] if kind == "MX" else [AAnswer()],
    )
    login(client)
    create_initial_member(client, username="missing.mailbox")

    assert vacation_app.deliver_outbox_once() is True
    assert vacation_app.deliver_outbox_once() is False
    with vacation_app.db() as conn:
        failed = conn.execute(
            "SELECT status, attempts, last_error FROM mail_outbox"
        ).fetchone()
    assert failed["status"] == "failed"
    assert failed["attempts"] == 1
    assert "Dauerhafte Ablehnung" in failed["last_error"]


def test_profile_picture_appears_in_navbar_and_matrix_but_not_export(client):
    login(client)
    response = client.post(
        "/profile",
        data={
            "email": "admin@example.de",
            "birth_date": "",
            "profile_image": (BytesIO(b"test-image"), "portrait.png"),
        },
        content_type="multipart/form-data",
        follow_redirects=True,
    )
    assert response.status_code == 200
    planner = client.get("/")
    assert planner.data.count(b"/profile-image/user-1.png") == 2
    assert b"avatar-nav" in planner.data
    assert b"avatar-matrix" in planner.data

    workbook_response = client.post(
        "/download",
        data={"start_year": date.today().year, "end_year": date.today().year},
    )
    workbook = load_workbook(BytesIO(workbook_response.data))
    assert workbook.active._images == []
    assert "urlaubsuebersicht_" in workbook_response.headers["Content-Disposition"]
    assert ".xlsx" in workbook_response.headers["Content-Disposition"]
    assert (
        workbook_response.mimetype
        == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    with ZipFile(BytesIO(workbook_response.data)) as archive:
        content_types = archive.read("[Content_Types].xml")
        relationships = archive.read("xl/_rels/workbook.xml.rels")
        archive_names = archive.namelist()
    assert (
        b"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"
        in content_types
    )
    assert b"application/vnd.ms-office.vbaProject" not in content_types
    assert b"relationships/vbaProject" not in relationships
    assert "xl/vbaProject.bin" not in archive_names

    deleted = client.post(
        "/profile",
        data={"action": "delete-profile-image"},
        follow_redirects=True,
    )
    assert b"Profilbild gel" in deleted.data
    assert b"/profile-image/user-1.png" not in client.get("/").data


def test_mobile_delete_control_and_responsive_matrix_are_rendered(client):
    login(client)
    planner = client.get("/")
    assert b'id="delete-selection"' in planner.data
    assert b"addEventListener('click',deleteSelection)" in planner.data
    assert b"matrix-selection.js" in planner.data
    selection_script = client.get("/static/matrix-selection.js").data
    assert b"selectRectangle" in selection_script
    assert b'cell.addEventListener("mouseenter"' in selection_script
    stylesheet = client.get("/static/style.css").data
    assert b"width:calc(100% - 2rem)" in stylesheet
    assert b"scroll-snap-type:x mandatory" in stylesheet
    assert b"border-right:4px solid #111" in stylesheet


def test_desksharing_docx_import_uses_iso_week_and_first_names(client, tmp_path):
    login(client)
    with vacation_app.db() as conn:
        anna_id = vacation_app.create_user(
            conn, "anna", "Anna", "Muster", "password", ["normal", "desksharing"]
        )
        ben_id = vacation_app.create_user(
            conn, "ben", "Ben", "Beispiel", "password", ["normal", "desksharing"]
        )
    document = Document()
    document.add_paragraph("KW 32")
    table = document.add_table(rows=3, cols=4)
    for column, value in enumerate(["Name", "Montag", "Dienstag", "Mittwoch"]):
        table.cell(0, column).text = value
    for column, value in enumerate(["Anna", "Anwesend", "Homeoffice", "Abwesend"]):
        table.cell(1, column).text = value
    for column, value in enumerate(["Ben", "Homeoffice", "Anwesend", ""]):
        table.cell(2, column).text = value
    path = tmp_path / "Desksharing-Detaillierung.docx"
    document.save(path)

    assert vacation_app.import_desksharing_word(path, 2026) == 5
    monday = date.fromisocalendar(2026, 32, 1).isoformat()
    tuesday = date.fromisocalendar(2026, 32, 2).isoformat()
    with vacation_app.db() as conn:
        assert (
            conn.execute(
                "SELECT status FROM desksharing_entries WHERE user_id = ? AND entry_date = ?",
                (anna_id, monday),
            ).fetchone()["status"]
            == "Anwesend"
        )
        assert (
            conn.execute(
                "SELECT status FROM desksharing_entries WHERE user_id = ? AND entry_date = ?",
                (ben_id, tuesday),
            ).fetchone()["status"]
            == "Anwesend"
        )


def test_legacy_word_text_import_and_admin_desksharing_edit(client, tmp_path):
    login(client)
    with vacation_app.db() as conn:
        user_id = vacation_app.create_user(
            conn, "lea", "Lea", "Muster", "password", ["normal", "desksharing"]
        )
        vacation_app.create_user(
            conn, "outsider", "Nicht", "Aufgelistet", "password", ["normal"]
        )
    path = tmp_path / "Desksharing-Detaillierung.doc"
    path.write_text("KW 33\nMontag\nHomeoffice: Lea\n", encoding="utf-8")
    assert vacation_app.import_desksharing_word(path, 2026) == 1
    monday = date.fromisocalendar(2026, 33, 1).isoformat()
    response = client.post(
        "/desksharing/bulk-entry",
        json={
            "status": "Anwesend",
            "cells": [{"user_id": user_id, "date": monday}],
        },
    )
    assert response.json == {"updated": 1}
    planner = client.get("/desksharing?year=2026")
    assert b"Desksharing 2026" in planner.data
    assert b"Anwesend" in planner.data
    assert b"Lea Muster" in planner.data
    assert b"Nicht Aufgelistet" not in planner.data

    client.get("/logout")
    login(client, "lea", "password")
    forbidden = client.post(
        "/desksharing/bulk-entry",
        json={"delete": True, "cells": [{"user_id": user_id, "date": monday}]},
    )
    assert forbidden.status_code == 403
    normal_view = client.get("/desksharing?year=2026")
    assert b'id="today"' in normal_view.data
    assert b"const desksharingPlanner" in normal_view.data
    assert (
        b"window.onload=()=>document.getElementById('today').click()"
        in normal_view.data
    )


def test_admin_can_persist_independent_matrix_orders(client):
    login(client)
    with vacation_app.db() as conn:
        first_id = vacation_app.create_user(
            conn, "first", "Erster", "Nutzer", "password", ["normal", "desksharing"]
        )
        second_id = vacation_app.create_user(
            conn,
            "second",
            "Zweiter",
            "Nutzer",
            "password",
            ["normal", "desksharing"],
        )
    vacation_order = [second_id, first_id, 1]
    saved = client.post("/matrix-order/vacation", json={"user_ids": vacation_order})
    assert saved.json == {"saved": 3}
    vacation_page = client.get("/").data
    vacation_rows = [
        vacation_page.index(f'<tr data-user-id="{user_id}"'.encode())
        for user_id in vacation_order
    ]
    assert vacation_rows == sorted(vacation_rows)

    desksharing_order = [first_id, second_id]
    saved = client.post(
        "/matrix-order/desksharing", json={"user_ids": desksharing_order}
    )
    assert saved.json == {"saved": 2}
    desksharing_page = client.get("/desksharing").data
    assert desksharing_page.index(b"Erster Nutzer") < desksharing_page.index(
        b"Zweiter Nutzer"
    )
    assert b"Sortiermodus starten" in desksharing_page
    assert b"Anordnung speichern" in desksharing_page
    assert b"matrix-order.js" in desksharing_page
    assert b"matrix-selection.js" in desksharing_page

    client.get("/logout")
    login(client, "first", "password")
    assert (
        client.post(
            "/matrix-order/desksharing", json={"user_ids": desksharing_order}
        ).status_code
        == 403
    )


def test_matrix_order_rejects_incomplete_user_lists(client):
    login(client)
    assert (
        client.post("/matrix-order/vacation", json={"user_ids": [1, 1]}).status_code
        == 400
    )


def test_admin_can_create_edit_and_protect_entry_mappings(client):
    login(client)
    page = client.get("/entry-mappings")
    assert page.status_code == 200
    assert b"Zuordnungen verwalten" in page.data
    assert b"UrlbGplntOdrBntrgt" in page.data

    mapping_data = {
        "import_code": "ImportTest",
        "matrix_color": "#112233",
        "matrix_code": "TX",
        "button_name": "Testzeit",
        "export_code": "TEX",
        "export_color": "#445566",
        "export_description": "Testzeit Export",
        "visible_roles": ["normal", "azubi"],
        "azubi__override": "1",
        "azubi__matrix_code": "AZTX",
        "azubi__matrix_color": "#123456",
        "azubi__export_code": "AZEX",
        "azubi__export_color": "#654321",
    }
    created = client.post("/entry-mappings", data=mapping_data, follow_redirects=True)
    assert b"Zuordnung gespeichert" in created.data
    with vacation_app.db() as conn:
        mapping = conn.execute(
            "SELECT * FROM entry_mappings WHERE matrix_code = 'TX'"
        ).fetchone()
        roles = vacation_app.mapping_roles(conn, mapping["id"])
    assert mapping["matrix_color"] == "112233"
    assert roles["normal"]["visible"] == 1
    assert roles["admin"]["visible"] == 0
    assert roles["azubi"]["matrix_code"] == "AZTX"
    assert roles["azubi"]["export_color"] == "654321"

    duplicate_data = dict(mapping_data, matrix_code="OTHER")
    duplicate = client.post(
        "/entry-mappings", data=duplicate_data, follow_redirects=True
    )
    assert b"bereits vergeben" in duplicate.data
    with vacation_app.db() as conn:
        assert (
            conn.execute(
                "SELECT COUNT(*) FROM entry_mappings WHERE import_code = 'ImportTest'"
            ).fetchone()[0]
            == 1
        )

    with vacation_app.db() as conn:
        conn.execute(
            "INSERT INTO entries(user_id, entry_date, code, created_by) VALUES (1, ?, 'TX', 1)",
            (date.today().isoformat(),),
        )
    edited_data = dict(mapping_data, mapping_id=mapping["id"], matrix_code="TY")
    edited = client.post("/entry-mappings", data=edited_data, follow_redirects=True)
    assert b"Zuordnung gespeichert" in edited.data
    with vacation_app.db() as conn:
        assert (
            conn.execute(
                "SELECT code FROM entries WHERE user_id = 1 AND entry_date = ?",
                (date.today().isoformat(),),
            ).fetchone()["code"]
            == "TY"
        )
        assert (
            conn.execute(
                "SELECT 1 FROM entry_mappings WHERE matrix_code = 'TX'"
            ).fetchone()
            is None
        )


def test_role_specific_mapping_controls_matrix_import_and_export(client, tmp_path):
    login(client)
    client.post(
        "/entry-mappings",
        data={
            "import_code": "RoleImport",
            "matrix_color": "#112233",
            "matrix_code": "RX",
            "button_name": "Rollenzeit",
            "export_code": "REX",
            "export_color": "#445566",
            "export_description": "Rollenzeit Export",
            "visible_roles": ["azubi"],
            "azubi__override": "1",
            "azubi__matrix_code": "AZRX",
            "azubi__matrix_color": "#123456",
            "azubi__export_code": "AZEX",
            "azubi__export_color": "#654321",
        },
    )
    with vacation_app.db() as conn:
        user_id = vacation_app.create_user(
            conn, "mapping.azubi", "Mapping", "Azubi", "password", ["azubi"]
        )

    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Name", date.today()])
    sheet.append(["Mapping Azubi", "RoleImport"])
    path = tmp_path / "role-import.xlsx"
    workbook.save(path)
    vacation_app.import_excel(path)
    with vacation_app.db() as conn:
        assert (
            conn.execute(
                "SELECT code FROM entries WHERE user_id = ?", (user_id,)
            ).fetchone()["code"]
            == "RX"
        )

    client.get("/logout")
    login(client, "mapping.azubi", "password")
    planner = client.get("/")
    assert b"AZRX" in planner.data
    assert b"--legend-color:#123456" in planner.data
    exported = client.post(
        "/download",
        data={"start_year": date.today().year, "end_year": date.today().year},
    )
    exported_workbook = load_workbook(BytesIO(exported.data))
    sheet = exported_workbook.active
    values = [cell.value for row in sheet.iter_rows() for cell in row]
    assert "AZEX" in values
    export_cell = next(
        cell for row in sheet.iter_rows() for cell in row if cell.value == "AZEX"
    )
    assert export_cell.fill.fgColor.rgb.endswith("654321")

    client.get("/logout")
    login(client)
    with vacation_app.db() as conn:
        vacation_app.create_user(
            conn, "mapping.normal", "Mapping", "Normal", "password", ["normal"]
        )
    client.get("/logout")
    login(client, "mapping.normal", "password")
    assert b"AZRX" not in client.get("/").data
    assert client.get("/entry-mappings").status_code == 403


def test_admin_can_delete_mapping_without_deleting_existing_entries(client):
    login(client)
    with vacation_app.db() as conn:
        mapping = conn.execute(
            "SELECT id FROM entry_mappings WHERE matrix_code = 'UB'"
        ).fetchone()
        conn.execute(
            "INSERT INTO entries(user_id, entry_date, code, created_by) VALUES (1, ?, 'UB', 1)",
            (date.today().isoformat(),),
        )

    response = client.post(
        f"/entry-mappings/{mapping['id']}/delete", follow_redirects=True
    )
    assert response.status_code == 200
    assert b"Zuordnung UB gel\xc3\xb6scht" in response.data
    with vacation_app.db() as conn:
        assert (
            conn.execute(
                "SELECT 1 FROM entry_mappings WHERE id = ?", (mapping["id"],)
            ).fetchone()
            is None
        )
        assert (
            conn.execute("SELECT code FROM entries WHERE code = 'UB'").fetchone()[
                "code"
            ]
            == "UB"
        )


def test_mapping_delete_is_admin_only_and_unknown_mapping_is_not_found(client):
    login(client)
    assert client.post("/entry-mappings/999999/delete").status_code == 404
    with vacation_app.db() as conn:
        mapping_id = conn.execute(
            "SELECT id FROM entry_mappings WHERE matrix_code = 'UB'"
        ).fetchone()["id"]
        vacation_app.create_user(
            conn, "mapping.normal.delete", "Mapping", "Normal", "password", ["normal"]
        )
    client.get("/logout")
    login(client, "mapping.normal.delete", "password")
    assert client.post(f"/entry-mappings/{mapping_id}/delete").status_code == 403
    with vacation_app.db() as conn:
        assert conn.execute(
            "SELECT 1 FROM entry_mappings WHERE id = ?", (mapping_id,)
        ).fetchone()


def test_repo_example_excel_imports_entries(client):
    login(client)
    example = (
        Path(__file__).resolve().parents[1] / "Urlaubsübersicht-Detaillierung.xlsx"
    )
    vacation_app.import_excel(example)
    with vacation_app.db() as conn:
        entry_count = conn.execute("SELECT COUNT(*) FROM entries").fetchone()[0]
        imported_user = conn.execute(
            "SELECT id FROM users WHERE username = ?", ("alexander.hälter",)
        ).fetchone()
    assert imported_user is not None
    assert entry_count > 100
